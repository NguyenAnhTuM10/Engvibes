# -*- coding: utf-8 -*-
"""Core document engine: styles, helpers, field codes for TOC / captions."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Times New Roman"
SIZE = 13


def _set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_repeat_header(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader")
    th.set(qn("w:val"), "true")
    trPr.append(th)


class Report:
    def __init__(self):
        self.doc = Document()
        self._setup_styles()
        self._setup_page()

    # ---------- setup ----------
    def _setup_page(self):
        sec = self.doc.sections[0]
        sec.page_height = Cm(29.7)
        sec.page_width = Cm(21.0)
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(3.0)
        sec.right_margin = Cm(2.0)

    def _setup_styles(self):
        styles = self.doc.styles
        normal = styles["Normal"]
        normal.font.name = FONT
        normal.font.size = Pt(SIZE)
        normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        pf = normal.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.5
        pf.space_after = Pt(6)
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # heading colors -> black, bold, Times New Roman
        for i, sz in [(1, 16), (2, 14), (3, 13), (4, 13)]:
            h = styles[f"Heading {i}"]
            h.font.name = FONT
            h.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
            h.font.size = Pt(sz)
            h.font.bold = True
            h.font.color.rgb = RGBColor(0, 0, 0)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)
            h.paragraph_format.line_spacing = 1.5
            h.paragraph_format.keep_with_next = True

        if "Caption" in [s.name for s in styles]:
            cap = styles["Caption"]
            cap.font.name = FONT
            cap.font.size = Pt(12)
            cap.font.italic = True
            cap.font.bold = False
            cap.font.color.rgb = RGBColor(0, 0, 0)
            cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_before = Pt(4)
            cap.paragraph_format.space_after = Pt(10)

    # ---------- text ----------
    def h1(self, text):
        self.doc.add_page_break()
        p = self.doc.add_heading(text, level=1)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return p

    def h1_nobreak(self, text):
        p = self.doc.add_heading(text, level=1)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return p

    def h2(self, text):
        return self.doc.add_heading(text, level=2)

    def h3(self, text):
        return self.doc.add_heading(text, level=3)

    def h4(self, text):
        return self.doc.add_heading(text, level=4)

    def p(self, text="", bold=False, italic=False, align=None, size=None):
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic
        if size:
            run.font.size = Pt(size)
        if align == "center":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "left":
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif align == "right":
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        return para

    def bullet(self, text, level=0):
        p = self.doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(1.0 + 0.6 * level)
        p.add_run(text)
        return p

    def num(self, text):
        p = self.doc.add_paragraph(style="List Number")
        p.add_run(text)
        return p

    def para_rich(self, segments):
        """segments: list of (text, bold, italic)"""
        para = self.doc.add_paragraph()
        for seg in segments:
            text = seg[0]
            bold = seg[1] if len(seg) > 1 else False
            italic = seg[2] if len(seg) > 2 else False
            r = para.add_run(text)
            r.bold = bold
            r.italic = italic
        return para

    def code(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(8)
        for line in text.split("\n"):
            r = p.add_run(line)
            r.font.name = "Consolas"
            r.font.size = Pt(10)
            r.add_break()
        return p

    # ---------- captions with SEQ fields ----------
    def figure(self, caption, placeholder=True, height_cm=6.5):
        if placeholder:
            box = self.doc.add_paragraph()
            box.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = box.add_run(f"[ Hình minh hoạ: {caption} ]")
            r.italic = True
            r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            box.paragraph_format.space_before = Pt(4)
        self._caption("Hình", caption)

    def table_caption(self, caption):
        self._caption("Bảng", caption)

    def _caption(self, label, text):
        p = self.doc.add_paragraph(style="Caption")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label} ")
        self._add_seq_field(p, label)
        p.add_run(f": {text}")

    def _add_seq_field(self, paragraph, label):
        run = paragraph.add_run()
        fldBegin = OxmlElement("w:fldChar")
        fldBegin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = f" SEQ {label} \\* ARABIC "
        fldEnd = OxmlElement("w:fldChar")
        fldEnd.set(qn("w:fldCharType"), "end")
        run._r.append(fldBegin)
        run._r.append(instr)
        run._r.append(fldEnd)

    # ---------- tables ----------
    def table(self, header, rows, widths=None, caption=None, header_above=True):
        if caption and header_above:
            self.table_caption(caption)
        ncols = len(header)
        t = self.doc.add_table(rows=1, cols=ncols)
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = t.rows[0].cells
        for i, htext in enumerate(header):
            hdr[i].text = ""
            para = hdr[i].paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = para.add_run(htext)
            r.bold = True
            r.font.size = Pt(12)
            _set_cell_bg(hdr[i], "D9E2F3")
        _set_repeat_header(t.rows[0])
        for row in rows:
            cells = t.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = ""
                para = cells[i].paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                r = para.add_run(str(val))
                r.font.size = Pt(12)
        if widths:
            for i, w in enumerate(widths):
                for row in t.rows:
                    row.cells[i].width = Cm(w)
        sp = self.doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(2)
        return t

    # ---------- field-based TOC / lists ----------
    def toc(self, switches):
        p = self.doc.add_paragraph()
        run = p.add_run()
        b = OxmlElement("w:fldChar")
        b.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = switches
        sep = OxmlElement("w:fldChar")
        sep.set(qn("w:fldCharType"), "separate")
        t = OxmlElement("w:r")
        tt = OxmlElement("w:t")
        tt.text = "Nhấn Ctrl+A rồi F9 trong Word để cập nhật mục lục này."
        t.append(tt)
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.append(b)
        run._r.append(instr)
        run._r.append(sep)
        run._r.append(t)
        run._r.append(end)

    def page_break(self):
        self.doc.add_page_break()

    def save(self, path):
        self.doc.save(path)
